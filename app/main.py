from pathlib import Path
import shutil
import re
import base64
import secrets
import hmac
import hashlib
import time
from datetime import datetime, timezone
from uuid import uuid4

import requests
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, RedirectResponse, Response
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from docx import Document
from openpyxl import load_workbook

from app.config import get_settings
from app.ingestion.loaders import iter_supported_files, load_file, SUPPORTED_EXTENSIONS
from app.ingestion.chunking import make_chunks
from app.rag.vector_store import FaissStore
from app.rag.models import QueryRequest, QueryResponse, SaveAnswerRequest
from app.rag.pipeline import answer_query
from app.rag.chat import answer_chat_question
from app.rag.chat_store import ChatStore
from app.rag.memory import delete_qa_from_disk, save_qa_to_disk, rebuild_memory_index, list_memory_items
from dotenv import load_dotenv

load_dotenv()


settings = get_settings()
DEFAULT_RFP_TEMPLATE_PATH = Path("app/templates/AORN LMS Evaluation.docx")
UPLOADED_RFP_TEMPLATE_DIR = Path(settings.data_dir) / "rfp_templates"
UPLOADED_RFP_TEMPLATE_STEM = "uploaded-rfp-template"
GENERATED_RFP_DIR = Path(settings.data_dir) / "generated_rfps"
chat_store = ChatStore(settings.chat_history_dir)
SUPPORTED_RFP_TEMPLATE_EXTENSIONS = {".docx", ".xlsx", ".xlsm"}

app = FastAPI(title=settings.app_name)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/static", StaticFiles(directory="app/static"), name="static")
templates = Jinja2Templates(directory="app/templates")
SESSION_COOKIE_NAME = "rfp_assistant_session"
SESSION_MAX_AGE_SECONDS = 60 * 60 * 12


class RfpQuestionItem(BaseModel):
    question: str = Field(min_length=1, max_length=4000)
    answer: str = Field(default="", max_length=12000)


class RfpBatchRequest(BaseModel):
    questions: list[str] = Field(min_length=1, max_length=200)
    use_memory: bool = True
    use_documents: bool = True


class RfpSingleAnswerRequest(BaseModel):
    question: str = Field(min_length=1, max_length=4000)
    use_memory: bool = True
    use_documents: bool = True


class RfpBuildRequest(BaseModel):
    items: list[RfpQuestionItem] = Field(min_length=1, max_length=200)
    title: str = Field(default="Generated RFP Response", max_length=200)


class ChatConversationRequest(BaseModel):
    title: str = Field(default="New chat", max_length=100)


class ChatMessageRequest(BaseModel):
    message: str = Field(min_length=1, max_length=4000)


def _auth_secret() -> str:
    return f"{settings.app_username}:{settings.app_password or ''}"


def _sign_session(username: str, issued_at: int) -> str:
    payload = f"{username}:{issued_at}"
    signature = hmac.new(
        _auth_secret().encode("utf-8"),
        payload.encode("utf-8"),
        hashlib.sha256,
    ).hexdigest()
    return f"{payload}:{signature}"


def _is_valid_session(token: str | None) -> bool:
    if not token or not settings.app_password:
        return False

    username, _, rest = token.partition(":")
    issued_at_value, _, signature = rest.partition(":")
    if not username or not issued_at_value or not signature:
        return False

    try:
        issued_at = int(issued_at_value)
    except ValueError:
        return False

    if int(time.time()) - issued_at > SESSION_MAX_AGE_SECONDS:
        return False

    expected = _sign_session(username, issued_at)
    return secrets.compare_digest(token, expected) and secrets.compare_digest(username, settings.app_username)


def _is_basic_authorized(request: Request) -> bool:
    if not settings.app_password:
        return True

    authorization = request.headers.get("Authorization", "")
    scheme, _, token = authorization.partition(" ")
    if scheme.lower() != "basic" or not token:
        return False

    try:
        decoded = base64.b64decode(token).decode("utf-8")
    except Exception:
        return False

    username, _, password = decoded.partition(":")
    return (
        secrets.compare_digest(username, settings.app_username)
        and secrets.compare_digest(password, settings.app_password)
    )


def _is_authorized(request: Request) -> bool:
    if not settings.app_password:
        return True

    return _is_valid_session(request.cookies.get(SESSION_COOKIE_NAME)) or _is_basic_authorized(request)


def _login_page(error: str = "") -> HTMLResponse:
    error_html = f'<div class="error">{error}</div>' if error else ""
    return HTMLResponse(f"""<!doctype html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>RFP Assistant Login</title>
    <style>
        body {{
            margin: 0;
            min-height: 100vh;
            display: grid;
            place-items: center;
            background: #f5f7fb;
            color: #172033;
            font-family: Arial, sans-serif;
        }}
        main {{
            width: min(380px, calc(100vw - 32px));
            background: #ffffff;
            border: 1px solid #d8deea;
            border-radius: 8px;
            padding: 28px;
            box-shadow: 0 16px 40px rgba(23, 32, 51, 0.12);
        }}
        h1 {{
            margin: 0 0 20px;
            font-size: 24px;
            line-height: 1.2;
        }}
        label {{
            display: block;
            margin: 14px 0 6px;
            font-size: 14px;
            font-weight: 700;
        }}
        input {{
            width: 100%;
            box-sizing: border-box;
            border: 1px solid #b8c2d6;
            border-radius: 6px;
            padding: 11px 12px;
            font-size: 16px;
        }}
        button {{
            width: 100%;
            margin-top: 20px;
            border: 0;
            border-radius: 6px;
            background: #2557d6;
            color: #ffffff;
            padding: 12px 14px;
            font-size: 16px;
            font-weight: 700;
            cursor: pointer;
        }}
        .error {{
            margin-bottom: 14px;
            border: 1px solid #f1a6a6;
            border-radius: 6px;
            background: #fff1f1;
            color: #9d1c1c;
            padding: 10px 12px;
            font-size: 14px;
        }}
    </style>
</head>
<body>
    <main>
        <h1>RFP Assistant</h1>
        {error_html}
        <form method="post" action="/login">
            <label for="username">Username</label>
            <input id="username" name="username" autocomplete="username" required>
            <label for="password">Password</label>
            <input id="password" name="password" type="password" autocomplete="current-password" required>
            <button type="submit">Sign in</button>
        </form>
    </main>
</body>
</html>""")


@app.middleware("http")
async def require_app_password(request: Request, call_next):
    public_paths = {"/health", "/login"}
    if request.url.path in public_paths or _is_authorized(request):
        return await call_next(request)

    if request.method == "GET" and "text/html" in request.headers.get("accept", ""):
        return RedirectResponse(url="/login", status_code=303)

    return Response(
        status_code=401,
        headers={"WWW-Authenticate": 'Basic realm="RFP Assistant"'},
    )


def _normalize_match_text(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", value.lower()).strip()


def _get_rfp_template_path() -> Path:
    for suffix in SUPPORTED_RFP_TEMPLATE_EXTENSIONS:
        candidate = UPLOADED_RFP_TEMPLATE_DIR / f"{UPLOADED_RFP_TEMPLATE_STEM}{suffix}"
        if candidate.exists():
            return candidate

    return DEFAULT_RFP_TEMPLATE_PATH


class RateLimitError(Exception):
    def __init__(self, retry_after: int, message: str):
        super().__init__(message)
        self.retry_after = retry_after


def _answer_with_retry(
    question: str,
    use_memory: bool = True,
    use_documents: bool = True,
    attempts: int = 2,
) -> QueryResponse:
    last_error: Exception | None = None

    for attempt in range(attempts):
        try:
            return answer_query(
                query=question,
                use_memory=use_memory,
                use_documents=use_documents,
            )
        except Exception as exc:
            last_error = exc
            if isinstance(exc, requests.HTTPError) and exc.response is not None:
                if exc.response.status_code == 429:
                    retry_after_header = exc.response.headers.get("Retry-After")
                    retry_after = int(retry_after_header) if retry_after_header and retry_after_header.isdigit() else 25
                    raise RateLimitError(
                        retry_after=min(max(retry_after, 10), 60),
                        message="LLM rate limit reached. The queue will pause and retry.",
                    ) from exc

            if attempt < attempts - 1:
                continue

    raise RuntimeError(str(last_error) if last_error else "Answer generation failed")


def _document_references(response: QueryResponse) -> list[dict[str, str | float | int | None]]:
    references = []
    seen: set[tuple[str, int | None]] = set()

    for match in response.document_matches:
        source = match.metadata.get("source") or match.metadata.get("path") or match.id
        chunk_index = match.metadata.get("chunk_index")
        key = (str(source), chunk_index if isinstance(chunk_index, int) else None)
        if key in seen:
            continue
        seen.add(key)
        references.append({
            "source": str(source),
            "chunk_index": chunk_index if isinstance(chunk_index, int) else None,
            "score": match.score,
        })

    return references


def _clean_cell_text(value) -> str:
    if value is None:
        return ""

    return str(value).replace("\r", "\n").strip()


def _looks_like_question(value: str) -> bool:
    text = " ".join(value.split())
    if len(text) < 8:
        return False

    lowered = text.lower()
    if lowered in {"question", "questions", "requirement", "requirements", "prompt", "prompts"}:
        return False

    starts = (
        "describe", "provide", "explain", "confirm", "detail", "list", "identify",
        "state", "does", "do ", "is ", "are ", "will ", "can ", "what ", "how ",
        "when ", "where ", "why ", "please ",
    )
    return "?" in text or lowered.startswith(starts)


def _append_unique_question(questions: list[str], seen: set[str], question: str) -> None:
    cleaned = " ".join(question.split()).strip()
    key = _normalize_match_text(cleaned)
    if len(key) < 8 or key in seen:
        return

    seen.add(key)
    questions.append(cleaned)


def _extract_docx_template_questions(template_path: Path) -> list[str]:
    if not template_path.exists():
        return []

    doc = Document(template_path)
    questions: list[str] = []
    seen: set[str] = set()

    for table in doc.tables:
        response_col = None
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if "vendor response" in cell.text.lower():
                    response_col = idx
                    break
            if response_col is not None:
                break

        if response_col is None:
            continue

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) <= response_col:
                continue

            source_cells = [
                value for idx, value in enumerate(cells)
                if idx != response_col and value and "vendor response" not in value.lower()
            ]
            if not source_cells:
                continue

            question = max(source_cells, key=len).strip()
            key = _normalize_match_text(question)
            if len(key) < 20 or key in seen:
                continue

            seen.add(key)
            questions.append(question)

    return questions


def _extract_excel_template_questions(template_path: Path) -> list[str]:
    workbook = load_workbook(template_path, read_only=True, data_only=True)
    questions: list[str] = []
    seen: set[str] = set()
    header_markers = ("question", "requirement", "prompt")

    for sheet in workbook.worksheets:
        question_columns: set[int] = set()

        for row in sheet.iter_rows(values_only=True):
            values = [_clean_cell_text(value) for value in row]
            non_empty_values = [value for value in values if value]
            if not non_empty_values:
                continue

            header_columns = {
                index
                for index, value in enumerate(values)
                if any(marker in value.lower() for marker in header_markers)
            }
            if header_columns:
                question_columns.update(header_columns)
                continue

            for index in sorted(question_columns):
                if index < len(values) and values[index]:
                    candidate = values[index]
                    if _looks_like_question(candidate) or len(_normalize_match_text(candidate)) >= 20:
                        _append_unique_question(questions, seen, candidate)

            for value in non_empty_values:
                if _looks_like_question(value):
                    _append_unique_question(questions, seen, value)

    workbook.close()
    return questions


def _extract_template_questions() -> list[str]:
    template_path = _get_rfp_template_path()
    if not template_path.exists():
        return []

    suffix = template_path.suffix.lower()
    if suffix == ".docx":
        return _extract_docx_template_questions(template_path)
    if suffix in {".xlsx", ".xlsm"}:
        return _extract_excel_template_questions(template_path)

    return []


def _fill_template_response_cells(doc: Document, answers_by_question: dict[str, str]) -> set[str]:
    normalized_answers = {
        _normalize_match_text(question): answer
        for question, answer in answers_by_question.items()
        if question.strip() and answer.strip()
    }
    used_question_keys: set[str] = set()

    for table in doc.tables:
        response_col = None
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if "vendor response" in cell.text.lower():
                    response_col = idx
                    break
            if response_col is not None:
                break

        if response_col is None:
            continue

        for row in table.rows:
            if len(row.cells) <= response_col:
                continue

            response_cell = row.cells[response_col]
            if response_cell.text.strip():
                continue

            row_text = _normalize_match_text(
                " ".join(
                    cell.text for idx, cell in enumerate(row.cells)
                    if idx != response_col
                )
            )

            for question_key, answer in normalized_answers.items():
                if len(question_key) < 20 or len(row_text) < 20:
                    continue

                if question_key in row_text or row_text in question_key:
                    response_cell.text = answer
                    used_question_keys.add(question_key)
                    break

    return used_question_keys


def _append_unmatched_response_paragraphs(
    doc: Document,
    title: str,
    items: list[RfpQuestionItem],
    used_question_keys: set[str],
) -> None:
    unmatched_items = [
        item for item in items
        if _normalize_match_text(item.question) not in used_question_keys
    ]

    if not unmatched_items:
        return

    doc.add_page_break()
    doc.add_heading(title, level=1)

    for index, item in enumerate(unmatched_items, start=1):
        doc.add_heading(f"Response {index}", level=2)
        question = doc.add_paragraph()
        question.add_run("Question: ").bold = True
        question.add_run(item.question)
        answer = doc.add_paragraph()
        answer.add_run("Answer: ").bold = True
        answer.add_run(item.answer or "I could not find sufficient information in the provided data.")


@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.get("/chat", response_class=HTMLResponse)
def chat_page(request: Request):
    return templates.TemplateResponse("chat.html", {"request": request})


@app.get("/login", response_class=HTMLResponse)
def login_form():
    if not settings.app_password:
        return RedirectResponse(url="/", status_code=303)

    return _login_page()


@app.post("/login")
def login(username: str = Form(...), password: str = Form(...)):
    if not settings.app_password:
        return RedirectResponse(url="/", status_code=303)

    username_ok = secrets.compare_digest(username, settings.app_username)
    password_ok = secrets.compare_digest(password, settings.app_password)
    if not username_ok or not password_ok:
        return _login_page("Invalid username or password.")

    response = RedirectResponse(url="/", status_code=303)
    response.set_cookie(
        key=SESSION_COOKIE_NAME,
        value=_sign_session(settings.app_username, int(time.time())),
        max_age=SESSION_MAX_AGE_SECONDS,
        httponly=True,
        samesite="lax",
    )
    return response


@app.post("/logout")
def logout():
    response = RedirectResponse(url="/login", status_code=303)
    response.delete_cookie(SESSION_COOKIE_NAME)
    return response


@app.get("/health")
def health():
    llm_model = settings.ollama_model
    if settings.llm_provider == "openai":
        llm_model = settings.openai_model

    return {
        "status": "ok",
        "llm_provider": settings.llm_provider,
        "llm_model": llm_model,
        "embedding_model": settings.embedding_model,
    }


@app.post("/upload")
def upload_document(file: UploadFile = File(...)):
    filename = Path(file.filename or "").name
    suffix = Path(filename).suffix.lower()

    if not filename:
        raise HTTPException(status_code=400, detail="Filename is required")

    if suffix not in SUPPORTED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {suffix}")

    destination = settings.documents_dir / filename

    with destination.open("wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    return {
        "success": True,
        "filename": filename,
        "path": str(destination),
    }


@app.post("/rfp/template/upload")
def upload_rfp_template(file: UploadFile = File(...)):
    filename = Path(file.filename or "").name
    suffix = Path(filename).suffix.lower()

    if not filename:
        raise HTTPException(status_code=400, detail="Filename is required")

    if suffix not in SUPPORTED_RFP_TEMPLATE_EXTENSIONS:
        raise HTTPException(status_code=400, detail="RFP template must be a .docx, .xlsx, or .xlsm file")

    UPLOADED_RFP_TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    for existing in UPLOADED_RFP_TEMPLATE_DIR.glob(f"{UPLOADED_RFP_TEMPLATE_STEM}.*"):
        existing.unlink()

    uploaded_template_path = UPLOADED_RFP_TEMPLATE_DIR / f"{UPLOADED_RFP_TEMPLATE_STEM}{suffix}"

    with uploaded_template_path.open("wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    return {
        "success": True,
        "filename": filename,
        "template": uploaded_template_path.name,
        "questions": _extract_template_questions(),
    }


@app.post("/ingest/documents")
def ingest_documents():
    all_chunks = []
    files = list(iter_supported_files(settings.documents_dir))

    for file_path in files:
        text_parts = load_file(file_path)
        chunks = make_chunks(file_path, text_parts)
        all_chunks.extend(chunks)

    store = FaissStore("documents", settings.vector_path)
    store.rebuild(all_chunks)

    return {
        "success": True,
        "files_indexed": len(files),
        "chunks_indexed": len(all_chunks),
    }


@app.post("/query", response_model=QueryResponse)
def query(request: QueryRequest):
    return answer_query(
        query=request.query,
        use_memory=request.use_memory,
        use_documents=request.use_documents,
    )


@app.get("/api/chat/conversations")
def list_chat_conversations():
    return {"conversations": chat_store.list()}


@app.post("/api/chat/conversations")
def create_chat_conversation(request: ChatConversationRequest):
    return {"conversation": chat_store.create(request.title)}


@app.get("/api/chat/conversations/{conversation_id}")
def get_chat_conversation(conversation_id: str):
    conversation = chat_store.get(conversation_id)
    if conversation is None:
        raise HTTPException(status_code=404, detail="Conversation not found")
    return {"conversation": conversation}


@app.delete("/api/chat/conversations/{conversation_id}")
def delete_chat_conversation(conversation_id: str):
    if not chat_store.delete(conversation_id):
        raise HTTPException(status_code=404, detail="Conversation not found")
    return {"success": True}


@app.post("/api/chat/conversations/{conversation_id}/messages")
def send_chat_message(conversation_id: str, request: ChatMessageRequest):
    conversation = chat_store.get(conversation_id)
    if conversation is None:
        raise HTTPException(status_code=404, detail="Conversation not found")

    question = request.message.strip()
    result = answer_chat_question(question, conversation.get("messages", []))
    updated = chat_store.add_exchange(
        conversation_id=conversation_id,
        question=question,
        answer=result["answer"],
        sources=result["sources"],
        in_scope=result["in_scope"],
    )
    return {
        "conversation": updated,
        "message": updated["messages"][-1],
    }


@app.get("/rfp/template/questions")
def rfp_template_questions():
    template_path = _get_rfp_template_path()
    return {
        "template": template_path.name,
        "questions": _extract_template_questions(),
    }


@app.post("/rfp/answer")
def answer_rfp_questions(request: RfpBatchRequest):
    answers = []

    for question in request.questions:
        cleaned_question = question.strip()
        if not cleaned_question:
            continue

        try:
            response = _answer_with_retry(
                question=cleaned_question,
                use_memory=request.use_memory,
                use_documents=request.use_documents,
            )
            answers.append({
                "question": cleaned_question,
                "answer": response.answer,
                "from_memory": response.from_memory,
                "response_tags": response.response_tags,
                "document_references": _document_references(response),
                "rate_limited": False,
                "retry_after": None,
                "error": None,
            })
        except RateLimitError as exc:
            answers.append({
                "question": cleaned_question,
                "answer": "",
                "from_memory": False,
                "response_tags": [],
                "document_references": [],
                "rate_limited": True,
                "retry_after": exc.retry_after,
                "error": str(exc),
            })
        except Exception as exc:
            answers.append({
                "question": cleaned_question,
                "answer": "",
                "from_memory": False,
                "response_tags": [],
                "document_references": [],
                "rate_limited": False,
                "retry_after": None,
                "error": str(exc),
            })

    return {
        "items": answers,
    }


@app.post("/rfp/answer-one")
def answer_one_rfp_question(request: RfpSingleAnswerRequest):
    try:
        response = _answer_with_retry(
            question=request.question.strip(),
            use_memory=request.use_memory,
            use_documents=request.use_documents,
        )
        return {
            "success": True,
            "question": request.question.strip(),
            "answer": response.answer,
            "from_memory": response.from_memory,
            "response_tags": response.response_tags,
            "document_references": _document_references(response),
            "rate_limited": False,
            "retry_after": None,
            "error": None,
        }
    except RateLimitError as exc:
        return {
            "success": False,
            "question": request.question.strip(),
            "answer": "",
            "from_memory": False,
            "response_tags": [],
            "document_references": [],
            "rate_limited": True,
            "retry_after": exc.retry_after,
            "error": str(exc),
        }
    except Exception as exc:
        return {
            "success": False,
            "question": request.question.strip(),
            "answer": "",
            "from_memory": False,
            "response_tags": [],
            "document_references": [],
            "rate_limited": False,
            "retry_after": None,
            "error": str(exc),
        }


@app.post("/rfp/generate")
def generate_rfp_document(request: RfpBuildRequest):
    selected_template_path = _get_rfp_template_path()
    template_path = selected_template_path if selected_template_path.suffix.lower() == ".docx" else DEFAULT_RFP_TEMPLATE_PATH
    if not template_path.exists():
        raise HTTPException(status_code=404, detail="RFP template not found")

    completed_items = [
        item for item in request.items
        if item.question.strip() and item.answer.strip()
    ]

    if not completed_items:
        raise HTTPException(status_code=400, detail="At least one answered question is required")

    GENERATED_RFP_DIR.mkdir(parents=True, exist_ok=True)
    document = Document(template_path)
    answers_by_question = {
        item.question: item.answer
        for item in completed_items
    }

    used_question_keys = _fill_template_response_cells(document, answers_by_question)
    _append_unmatched_response_paragraphs(document, request.title, completed_items, used_question_keys)

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")
    output_path = GENERATED_RFP_DIR / f"rfp-response-{timestamp}-{uuid4().hex[:8]}.docx"
    document.save(output_path)

    return FileResponse(
        output_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="rfp-response.docx",
    )


@app.post("/memory/save")
def save_answer(request: SaveAnswerRequest):
    item = save_qa_to_disk(request)
    total = rebuild_memory_index()

    return {
        "success": True,
        "saved": item,
        "memory_items_indexed": total,
    }


@app.post("/memory/rebuild")
def rebuild_memory():
    total = rebuild_memory_index()

    return {
        "success": True,
        "memory_items_indexed": total,
    }


@app.get("/memory/list")
def list_memory():
    return {
        "items": list_memory_items()
    }


@app.delete("/memory/{qa_id}")
def delete_memory_item(qa_id: str):
    deleted = delete_qa_from_disk(qa_id)

    if not deleted:
        raise HTTPException(status_code=404, detail="Saved answer not found")

    total = rebuild_memory_index()

    return {
        "success": True,
        "deleted_id": qa_id,
        "memory_items_indexed": total,
    }
